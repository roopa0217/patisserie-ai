"""
Builds evals/solution_doc.docx -- a Word document matching the course's
"Solution Doc" template (The Primer + The Framework), ready to upload to
Google Drive and open with Google Docs (auto-converts, same workflow as the
xlsx -> Sheets import).

Usage:
    python -m scripts.build_solution_doc
"""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor

EVALS_DIR = Path(__file__).resolve().parent.parent.parent / "evals"
DIAGRAM_PATH = EVALS_DIR / "architecture.png"
OUT_PATH = EVALS_DIR / "solution_doc.docx"

NAVY = RGBColor(0x1F, 0x29, 0x37)
GRAY = RGBColor(0x55, 0x55, 0x55)


def add_heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = NAVY
    return h


def add_code_block(doc, code: str):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(8)
    run = p.add_run(code)
    run.font.name = "Courier New"
    run.font.size = Pt(9)
    run.font.color.rgb = GRAY
    return p


def add_table(doc, headers: list[str], rows: list[list[str]], col_widths=None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Light Grid Accent 1"
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = h
        for p in hdr_cells[i].paragraphs:
            for r in p.runs:
                r.font.bold = True
    for row in rows:
        cells = table.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = str(val)
    if col_widths:
        for row in table.rows:
            for i, w in enumerate(col_widths):
                row.cells[i].width = Inches(w)
    doc.add_paragraph()
    return table


def main():
    doc = Document()

    # ── Base style ────────────────────────────────────────────────────────────
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(10.5)

    # ── Title ─────────────────────────────────────────────────────────────────
    title = doc.add_heading("Patisserie AI — Intent Classifier Evaluation", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub = doc.add_paragraph("Solution Doc")
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub.runs[0].font.size = Pt(14)
    sub.runs[0].font.color.rgb = GRAY
    doc.add_paragraph()

    # ── The Primer ────────────────────────────────────────────────────────────
    add_heading(doc, "The Primer (evaluation one-liner)", level=1)
    primer = doc.add_paragraph()
    run = primer.add_run(
        "I will measure exact-match accuracy and LLM-judge agreement on Patisserie "
        "AI's intent classifier (_classify_intent in app/agent/graph.py) using a "
        "golden dataset of 50 cases covering happy-path, edge-case, known-failure, "
        "and adversarial instructor queries, with code-based exact-match plus a "
        "categorical LLM-as-judge. Pass bar: 90% exact-match accuracy with zero "
        "regressions. I will run this in LangSmith and report the delta from a 64% "
        "baseline to post-improvement."
    )
    run.italic = True

    # ── The Framework ─────────────────────────────────────────────────────────
    add_heading(doc, "The Framework", level=1)
    framework_rows = [
        (
            "Agent under test",
            "Patisserie AI's intent classifier (_classify_intent(), backend/app/agent/graph.py) "
            "— routes every instructor message to one of 5 categories before any tool runs: "
            "find_recipe, scale_recipe, build_indent, check_anomaly, general. Keyword-first, "
            "LLM fallback for anything that matches no keyword. See architecture diagram below.",
        ),
        (
            "User outcome",
            "The instructor gets routed to the right tool on the first try. A wrong route either "
            "blocks a useful answer or triggers the wrong action entirely.",
        ),
        (
            "Metrics (3-5)",
            "(1) Exact-match accuracy vs. ground truth. (2) Per-category precision/recall/F1. "
            "(3) LLM-judge agreement with human failure-category labeling. (4) Regression count "
            "vs. prior passing cases.",
        ),
        (
            "Judge method",
            "Exact match (code-based, free) for the primary metric. Categorical LLM-as-judge "
            "(sees query + predicted intent, never ground truth) for judge-alignment analysis. "
            "A second, independent general-correctness LLM-judge via the LangSmith SDK.",
        ),
        (
            "Golden dataset",
            "50 cases, hand-written: 25 happy / 15 edge / 7 known-failure / 3 adversarial. Every "
            "row traced to a specific keyword or regex path in the actual code. Stored in "
            "evals/golden_dataset.csv; also uploaded to LangSmith as 'patisserie-intent-classifier'.",
        ),
        (
            "Pass bar",
            "90% exact-match accuracy, zero regressions against the 16 pre-existing cases in "
            "backend/scripts/evaluate.py.",
        ),
        (
            "Instrumentation",
            "LangSmith traces the real _classify_intent() function per row via the SDK -- input, "
            "predicted output, exact-match score, LLM-judge score + reasoning per run.",
        ),
        (
            "Baseline run",
            "32/50 (64%). LangSmith experiment: intent-classifier-32d36008 (project "
            "patisserie-intent-eval).",
        ),
        (
            "Failure analysis",
            "18 failures, 4 root causes: (1) domain keyword bleed [12] -- a general question "
            "misrouted by an incidental keyword; (2) scale_recipe wins by dict-order priority [3] "
            "-- e.g. 'double check' hijacked by 'double'; (3) 'portions' hijacks recipe-lookup "
            "questions [2]; (4) 'list...ingredients' regex over-triggers build_indent [1].",
        ),
        (
            "Improvement hypotheses (3-4)",
            "(1) Gate 8 ambiguous keywords behind a WH-question check, predicted +12-16%. "
            "(2) Tighten the LLM fallback's own prompt, predicted +4-6%. (3) Narrow the "
            "compile/list regex, predicted +2%. (4) Exclude 'double' from the 'double check' "
            "idiom, predicted +2%.",
        ),
        (
            "Post-improvement run",
            "43/50 (86%), +22 points, zero regressions -- verified against all 50 rows and all "
            "16 original evaluate.py cases. LangSmith experiment: intent-classifier-cd86cca9. "
            "intent_exact_match 86%, intent_llm_judge 84%.",
        ),
        (
            "What is next",
            "Top remaining failure: 'portions' vs. recipe-lookup ambiguity (t022, t026). "
            "Deliberately not patched -- a 2-example rule would be overfitting, not improvement. "
            "Monitoring strategy: track each category's share of traffic, alert if 'general' "
            "drops sharply (signature of keyword-bleed regressions creeping back).",
        ),
    ]
    add_table(doc, ["Field", "Fill in"], framework_rows, col_widths=[1.6, 5.2])

    # ── Architecture diagram ──────────────────────────────────────────────────
    add_heading(doc, "Agent Architecture", level=1)
    if DIAGRAM_PATH.exists():
        doc.add_picture(str(DIAGRAM_PATH), width=Inches(6.0))
    else:
        doc.add_paragraph("(architecture.png not found -- re-run scripts/build_solution_doc.py)")
    doc.add_paragraph()

    doc.add_page_break()

    # ── Appendix A ────────────────────────────────────────────────────────────
    add_heading(doc, "Appendix A — The 4 Fixes (code, before/after)", level=1)
    doc.add_paragraph(
        "All 4 changes are to the real agent code (app/agent/graph.py, app/agent/prompts.py) "
        "-- nothing in the eval scripts or the golden dataset was touched. Each was implemented, "
        "tested against its target rows, then verified against the full 50-row set with zero "
        "regressions before moving to the next."
    )
    fix_rows = [
        (
            "1",
            "Domain keyword bleeds into the wrong category",
            "graph.py",
            "Added _WEAK_KEYWORDS + _is_general_question(): 8 ambiguous keywords are skipped on "
            "WH-question phrasing, falling through to the LLM fallback.",
            "t028, t033, t034, t046, t048, t049",
        ),
        (
            "2",
            "Same cluster -- LLM fallback's own bias",
            "prompts.py",
            "Tightened INTENT_SYSTEM: check_anomaly now requires a named recipe; general "
            "explicitly covers conceptual ratio/technique questions.",
            "t025, t030, t031",
        ),
        (
            "3",
            "'list...ingredients' override over-triggers build_indent",
            "graph.py",
            "compile/list regex now requires 'compile', or 'list' together with 'all'.",
            "t027",
        ),
        (
            "4",
            "scale_recipe wins by dict-order priority (partial)",
            "graph.py",
            "'double' no longer matches scale_recipe when part of the 'double check' idiom.",
            "t032",
        ),
    ]
    add_table(doc, ["#", "Cluster", "File", "Fix", "Rows fixed"], fix_rows, col_widths=[0.3, 1.6, 0.8, 3.2, 1.1])

    add_code_block(
        doc,
        "# Fix 1 -- app/agent/graph.py\n"
        "_WEAK_KEYWORDS = {\n"
        '    "check_anomaly": {"ratio", "correct", "check", "review"},\n'
        '    "build_indent": {"order", "production day"},\n'
        '    "find_recipe": {"recipe", "technique", "how do"},\n'
        '    "scale_recipe": {"yield"},\n'
        "}\n"
        '_QUESTION_START = re.compile(r"^(what|why|how|does|is|are|can|could|would)\\b")\n\n'
        "def _is_general_question(low: str) -> bool:\n"
        "    return bool(_QUESTION_START.match(low.strip()))\n\n"
        "is_question = _is_general_question(low)\n"
        'is_double_check_idiom = bool(re.search(r"\\bdouble[\\s-]?check", low))  # Fix 4\n'
        "for intent, keywords in _INTENT_KEYWORDS.items():\n"
        "    weak = _WEAK_KEYWORDS.get(intent, set())\n"
        "    for kw in keywords:\n"
        '        if kw == "double" and is_double_check_idiom:\n'
        "            continue\n"
        "        if kw in low and not (kw in weak and is_question):\n"
        "            return intent",
    )
    add_code_block(
        doc,
        "# Fix 2 -- app/agent/prompts.py, INTENT_SYSTEM\n"
        "- check_anomaly: the message names a SPECIFIC recipe already in the system and asks to\n"
        "  check, validate, or find errors in THAT recipe's measurements against baking thresholds\n"
        "- general: ...including questions about what a correct/good/ideal ratio or method is in\n"
        '  general (not about a specific named recipe to validate), even if words like "ratio",\n'
        '  "correct", or "check" appear',
    )
    add_code_block(
        doc,
        "# Fix 3 -- app/agent/graph.py\n"
        "if re.search(r'\\bcompile\\b.{0,60}\\bingredients?\\b', low):\n"
        '    return "build_indent"\n'
        "if re.search(r'\\blist\\b.{0,60}\\ball\\b.{0,60}\\bingredients?\\b', low):\n"
        '    return "build_indent"',
    )
    doc.add_paragraph(
        "Why we stopped at 4, deliberately: a 5th candidate fix (special-casing 'portions' when "
        "it co-occurs with a recipe-lookup word, to fix t022/t026) was reverse-engineered from "
        "exactly those 2 failing sentences rather than a general linguistic pattern -- that's "
        "overfitting to the golden set, not a real improvement. Left unfixed on purpose. Same "
        "reasoning for the 3 remaining adversarial rows (t035, t036, t050)."
    ).italic = True

    doc.add_paragraph("[Screenshot placeholder: paste a screenshot of the Step 4 · Fix + Delta tab here.]").italic = True

    # ── Appendix B ────────────────────────────────────────────────────────────
    add_heading(doc, "Appendix B — LLM-as-Judge (local, Steps 5/6)", level=1)
    doc.add_paragraph(
        "Categorical judge -- sees only the query + predicted intent, never ground truth -- "
        "asked whether each remaining failure is an instance of the largest cluster "
        "('domain keyword bleed'):"
    )
    add_code_block(
        doc,
        "You are evaluating an intent classifier for a pastry-academy assistant.\n\n"
        "Failure category: {title}\n"
        "Definition: {definition}\n\n"
        "Given the User Input and the classifier's Predicted Intent, decide whether the\n"
        "classifier's mistake is an instance of this failure category.\n"
        "Return TRUE only when the query's real intent falls outside the Predicted\n"
        "Intent specifically because of the mechanism described above.\n"
        "Return FALSE for correct predictions and for all other kinds of mistakes.\n\n"
        "Return structured output with fields: label (TRUE or FALSE), reasoning (one sentence).",
    )
    doc.add_paragraph(
        "Judge/human agreement: 42.9% (3/7) on the remaining 7 failures post-fixes. The "
        "disagreements are legitimate, not noise: with most 'clean' keyword-bleed cases already "
        "resolved, what's left are the genuinely ambiguous/adversarial ones."
    )
    doc.add_paragraph("[Screenshot placeholder: paste a screenshot of the Step 5/6 tabs here.]").italic = True

    # ── Appendix C ────────────────────────────────────────────────────────────
    add_heading(doc, "Appendix C — LangSmith (separate SDK-based evaluation)", level=1)
    for bullet in [
        "Dataset: patisserie-intent-classifier (50 examples, columns query/intent/id)",
        "Target: the real _classify_intent() function, traced via the SDK -- not a "
        "Playground-recreated prompt, since most queries never reach an LLM at all",
        "Evaluators: free code-based Exact Match, plus a general LLM-as-judge (sees query + "
        "predicted intent, decides correct/incorrect, blind to ground truth)",
        "Project: patisserie-intent-eval -- smith.langchain.com -> Datasets & Experiments -> "
        "patisserie-intent-classifier",
        "Result after all 4 fixes: intent_exact_match 43/50 (86%), intent_llm_judge 42/50 (84%) "
        "-- the two metrics measure different things, so a small gap is expected, not a bug.",
    ]:
        doc.add_paragraph(bullet, style="List Bullet")
    doc.add_paragraph("[Screenshot placeholder: paste a screenshot of the LangSmith experiment compare view here.]").italic = True

    doc.save(OUT_PATH)
    print(f"Wrote {OUT_PATH}")


if __name__ == "__main__":
    main()
